import os
import re
from docx import Document
from parser_utils import parse_sections, extract_table, strip_table_from_text
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
 
TITLE_COLOR  = RGBColor(0, 32, 91)
BODY_COLOR   = RGBColor(0, 0, 0)
FOOTER_COLOR = RGBColor(0xE4, 0x00, 0x2B)
 
def _apply_theme(org_name):
    global TITLE_COLOR, FOOTER_COLOR
    ol = str(org_name).lower()
    if "bharat petroleum" in ol or "bpcl" in ol:
        TITLE_COLOR = RGBColor(0, 0x72, 0xCE)
        FOOTER_COLOR = RGBColor(0xFF, 0xC7, 0x2C)
    elif "indian oil" in ol or "iocl" in ol or "indianoil" in ol:
        TITLE_COLOR = RGBColor(0xFF, 0x66, 0x00)
        FOOTER_COLOR = RGBColor(0x00, 0x33, 0xAA)
    else:
        TITLE_COLOR = RGBColor(0, 32, 91)
        FOOTER_COLOR = RGBColor(0xE4, 0, 0x2B)
 
# Regex to detect raw chart placeholders the LLM sometimes emits
_CHART_LINE = re.compile(r'^\s*\[CHART:', re.IGNORECASE)
 
_OFFICIAL_DOC_TYPES = {
    "file_note", "office_memorandum", "office_notice",
    "circular", "purchase_note", "office_order"
}
 
 
def parse_official_header(body_text, doc_type=None, user_prompt=None):
    """
    Extracts the header fields (Ref, Date, Subject, Dept, To, From etc.)
    from the top of the LLM output and returns them as a dict,
    along with the remaining body text.

    FIX: The footer scanner (Prepared by / Approved by etc.) now only
    activates AFTER the last ## section has ended. This prevents it from
    accidentally stripping approval lines that appear INSIDE Section 7
    of a Purchase Note or at the end of a File Note body.
    """
    header_fields = {}
    lines = body_text.splitlines()
    body_lines = []
    in_header = True
    non_empty_count = 0
 
    for line in lines:
        sline = line.strip()
        if in_header:
            if sline:
                non_empty_count += 1
 
            if sline.startswith('##') or re.match(r'^[=\-_*]{3,}$', sline) or non_empty_count > 15:
                in_header = False
                if sline.startswith('##'):
                    body_lines.append(line)
                continue
 
            if sline.startswith('#'):
                header_fields['title'] = sline.lstrip('#').strip()
            elif ':' in sline:
                parts = sline.split(':', 1)
                k_candidate = parts[0].strip()
                v_candidate = parts[1].strip()
                k_norm = k_candidate.lower().replace(' ', '_').replace('.', '')
                if k_norm in {
                    'ref', 'ref_no', 'file_no', 'no', 'note_no', 'notice_no',
                    'order_no', 'circular_no', 'date', 'subject', 'sub', 'to',
                    'from', 'through', 'department', 'dept', 'title',
                    'reference', 'reference_no', 'organization', 'company', 'org'
                }:
                    header_fields[k_norm] = v_candidate
                else:
                    in_header = False
                    body_lines.append(line)
            elif sline != '':
                in_header = False
                body_lines.append(line)
        else:
            body_lines.append(line)
 
    if not body_lines:
        body_lines = lines
 
    full_text = "\n".join(lines[:15])
 
    # Extract ref/file number
    for pattern in [
        r'(?:File\s+No\.|Note\s+No\.|Notice\s+No\.|Circular\s+No\.|Order\s+No\.|No\.)\s*:\s*([^\n\t]+)',
        r'(?:Circular\s+No\.|Order\s+No\.)\s*([^\n\t]+)'
    ]:
        m = re.search(pattern, full_text, re.IGNORECASE)
        if m:
            val = m.group(1).strip()
            if 'date' in val.lower():
                val = re.split(r'\bdate\b', val, flags=re.IGNORECASE)[0].strip().rstrip(':').strip()
            header_fields['ref_no'] = val
            break
 
    m_date = re.search(r'Date\s*:\s*([^\n\t\r]+)', full_text, re.IGNORECASE)
    if m_date:
        header_fields['date'] = m_date.group(1).strip()
 
    m_subj = re.search(r'Subject\s*:\s*([^\n\t\r]+)', full_text, re.IGNORECASE)
    if m_subj:
        header_fields['subject'] = m_subj.group(1).strip()
 
    m_dept = re.search(r'Department\s*:\s*([^\n\t\r]+)', full_text, re.IGNORECASE)
    if m_dept:
        header_fields['department'] = m_dept.group(1).strip()
 
    m_to = re.search(r'\bTo\s*:\s*([^\n\t\r]+)', full_text, re.IGNORECASE)
    if m_to:
        header_fields['to'] = m_to.group(1).strip()
 
    m_from = re.search(r'\bFrom\s*:\s*([^\n\t\r]+)', full_text, re.IGNORECASE)
    if m_from:
        header_fields['from'] = m_from.group(1).strip()
 
    # Clean bracketed placeholders
    for k in ['ref_no', 'date', 'subject', 'department', 'to', 'from']:
        if k in header_fields and ('[' in header_fields[k] or ']' in header_fields[k]):
            header_fields[k] = ""
            
    # Scrub invented dates
    if 'date' in header_fields and header_fields['date']:
        date_val = header_fields['date']
        if '_' not in date_val and user_prompt:
            prompt_lower = user_prompt.lower()
            words = re.findall(r'\b\w+\b', date_val.lower())
            found = False
            for w in words:
                if w in prompt_lower:
                    found = True
                    break
            if not found:
                header_fields['date'] = "__________________"
 
    # -- FIX: Footer scanner --
    # Only scan for trailing signature/distribution lines AFTER the last ## section.
    # For Purchase Notes and File Notes, approval lines appear inside body sections
    # (e.g. ## 7. Recommendation) and must NOT be stripped from the body.
    # We detect "trailing" footer lines as those appearing AFTER the last ## heading
    # AND after all numbered section content has ended.
 
    footer_keys = {
        'prepared_by', 'recommended_by', 'approved_by',
        'submitted_for_approval_of', 'submitted_for_approval_to',
        'submitted_for_approval', 'submitted_by', 'signatory',
        'by_order_of', 'copy_to', 'distribution',
        'designation', 'sd/-',
    }
 
    # Find the index of the last ## heading in body_lines
    last_section_idx = -1
    for i, line in enumerate(body_lines):
        if line.strip().startswith('##'):
            last_section_idx = i
 
    # Only run footer scan after last section + reasonable content gap
    # For purchase_note specifically, skip the footer scan entirely --
    # the approval chain is part of Section 7 body content.
    skip_footer_scan = (doc_type == "purchase_note")
 
    final_body_lines = []
    if skip_footer_scan or last_section_idx == -1:
        # No footer scan: keep all body lines as-is
        final_body_lines = body_lines
    else:
        # Only scan lines that appear to be TRAILING (after last section + 20 lines buffer)
        # This is conservative -- we'd rather keep a line in body than strip it wrongly
        footer_scan_start = last_section_idx + 20
 
        for i, line in enumerate(body_lines):
            sline = line.strip()
            if i >= footer_scan_start and ':' in sline and not sline.startswith('##') and not sline.startswith('#'):
                parts = sline.split(':', 1)
                k_candidate = parts[0].strip()
                v_candidate = parts[1].strip()
                k_norm = k_candidate.lower().replace(' ', '_').replace('.', '').replace(' ', '_')
 
                if k_norm in footer_keys:
                    if k_norm == 'submitted_for_approval_of' and 'approved_by' not in header_fields:
                        header_fields['approved_by'] = v_candidate
                    elif k_norm == 'submitted_by' and 'prepared_by' not in header_fields:
                        header_fields['prepared_by'] = v_candidate
                    elif k_norm == 'sd/-':
                        header_fields['signatory'] = v_candidate
                    else:
                        header_fields[k_norm] = v_candidate
                    continue
            final_body_lines.append(line)
 
    if not final_body_lines:
        final_body_lines = body_lines
 
    return header_fields, "\n".join(final_body_lines).strip()
 
 
def _add_official_header(doc, metadata, doc_type):
    org_candidate = (metadata.get("organization") or metadata.get("company") or
                     metadata.get("org"))
    ref_candidate = (metadata.get("ref_no") or metadata.get("file_no") or
                     metadata.get("note_no") or metadata.get("no"))
 
    if not org_candidate:
        org_candidate = "__________________________________________________"
 
    org = org_candidate
    ref = ref_candidate or "__________________"
 
    # Company Name
    p_company = doc.add_paragraph()
    r_company = p_company.add_run(org.upper())
    r_company.font.size = Pt(14)
    r_company.font.bold = True
    r_company.font.color.rgb = TITLE_COLOR
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_company.space_after = Pt(2)
 
    # Department
    dept = metadata.get("department") or metadata.get("dept") or "Corporate Office"
    p_dept = doc.add_paragraph()
    r_dept = p_dept.add_run(dept)
    r_dept.font.size = Pt(11)
    r_dept.font.bold = True
    r_dept.font.color.rgb = BODY_COLOR
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dept.space_after = Pt(8)
 
    # Divider
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_div = p_div.add_run("―" * 50)
    r_div.font.color.rgb = RGBColor(180, 180, 180)
    p_div.space_after = Pt(6)
 
    # Ref and Date
    date = metadata.get("date") or "__________________"
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
 
    cell_l = tbl.cell(0, 0)
    p_l = cell_l.paragraphs[0]
    r_l = p_l.add_run(f"Ref: {ref}")
    r_l.font.bold = True
    r_l.font.size = Pt(10)
 
    cell_r = tbl.cell(0, 1)
    p_r = cell_r.paragraphs[0]
    r_r = p_r.add_run(f"Date: {date}")
    r_r.font.bold = True
    r_r.font.size = Pt(10)
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
 
    # Second divider
    p_div2 = doc.add_paragraph()
    p_div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_div2 = p_div2.add_run("―" * 50)
    r_div2.font.color.rgb = RGBColor(180, 180, 180)
    p_div2.space_before = Pt(6)
    p_div2.space_after = Pt(12)
 
    # Document Type Title
    doc_title_map = {
        "file_note":          "FILE NOTE",
        "office_memorandum":  "OFFICE MEMORANDUM",
        "office_notice":      "OFFICE NOTICE",
        "circular":           "CIRCULAR",
        "purchase_note":      "PURCHASE NOTE",
        "office_order":       "OFFICE ORDER"
    }
    dtype_name = doc_title_map.get(doc_type, doc_type.upper().replace("_", " "))
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_type = p_type.add_run(dtype_name)
    r_type.font.size = Pt(14)
    r_type.font.bold = True
    r_type.font.color.rgb = TITLE_COLOR
    p_type.space_after = Pt(12)
 
    # Subject
    subj = metadata.get("subject") or "Official Communication"
    p_subj = doc.add_paragraph()
    r_subj = p_subj.add_run(f"Subject: {subj}")
    r_subj.font.bold = True
    r_subj.font.size = Pt(11)
    r_subj.font.underline = True
    p_subj.space_after = Pt(12)
 
    # To / From / Through routing blocks
    has_routing = False
    route_tbl = doc.add_table(rows=0, cols=2)
    route_tbl.autofit = False
 
    for f in ["to", "through", "from"]:
        val = metadata.get(f)
        if val:
            has_routing = True
            row = route_tbl.add_row()
            p0 = row.cells[0].paragraphs[0]
            r0 = p0.add_run(f"{f.capitalize()}:")
            r0.font.bold = True
            r0.font.size = Pt(10)
            p1 = row.cells[1].paragraphs[0]
            r1 = p1.add_run(val)
            r1.font.size = Pt(10)
 
    if has_routing:
        p_space = doc.add_paragraph()
        p_space.space_after = Pt(12)
 
 
def _add_official_footer(doc, metadata):
    # For purchase_note, the approval chain is rendered as part of the body content.
    # Only render a separate footer block for other official doc types.
    sig_drawn = False
    for k, label in [("prepared_by", "Prepared by"), ("recommended_by", "Recommended by"),
                      ("approved_by", "Approved by")]:
        v = metadata.get(k)
        if v:
            sig_drawn = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(2)
            r_lbl = p.add_run(f"({label})")
            r_lbl.font.bold = True
            r_lbl.font.size = Pt(10)
 
            p_val = doc.add_paragraph()
            p_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_val.paragraph_format.space_before = Pt(0)
            p_val.paragraph_format.space_after = Pt(6)
            r_v = p_val.add_run(v)
            r_v.font.bold = True
            r_v.font.size = Pt(10)
 
    if not sig_drawn:
        v_sig = metadata.get("signatory") or metadata.get("by_order_of")
        if v_sig:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(v_sig)
            r.font.bold = True
            r.font.size = Pt(10)
 
    for k, label in [("copy_to", "Copy to:"), ("distribution", "Distribution:")]:
        val = metadata.get(k)
        if val:
            p_cop = doc.add_paragraph()
            p_cop.space_before = Pt(12)
            r_cop_hdr = p_cop.add_run(f"{label}\n")
            r_cop_hdr.font.bold = True
            r_cop_hdr.font.size = Pt(10)
            for item in val.split(','):
                it = item.strip()
                if it:
                    r_item = p_cop.add_run(f"  - {it}\n")
                    r_item.font.size = Pt(10)
 
 
def generate_docx(body, output_path, title="Generated Document", chart_images=None, doc_type=None, user_prompt=None):
    chart_images = _dedupe_charts(list(chart_images or []))
    doc = Document()
 
    body = _strip_chart_placeholders(body)
    body = _clean_bracketed_placeholders(body)
 
    is_official = doc_type in _OFFICIAL_DOC_TYPES
    metadata = {}
 
    org_name = ""
    if is_official:
        # Pass doc_type so footer scanner knows not to strip purchase_note body lines
        metadata, body = parse_official_header(body, doc_type=doc_type, user_prompt=user_prompt)
        org_name = (metadata.get('organization') or metadata.get('company') or
                    metadata.get('org') or "")
        if not org_name:
            bl = body.lower()
            if "bpcl" in bl or "bharat petroleum" in bl:
                org_name = "BHARAT PETROLEUM CORPORATION LIMITED"
            elif "iocl" in bl or "indian oil" in bl or "indianoil" in bl:
                org_name = "INDIAN OIL CORPORATION LIMITED"
            elif "ongc" in bl or "oil and natural gas" in bl:
                org_name = "OIL AND NATURAL GAS CORPORATION LIMITED"
            elif "gail" in bl:
                org_name = "GAIL (INDIA) LIMITED"
            elif "hpcl" in bl or "hindustan petroleum" in bl:
                org_name = "HINDUSTAN PETROLEUM CORPORATION LIMITED"
            else:
                org_name = "__________________________________________________"
        metadata['organization'] = org_name
        _apply_theme(org_name)
        _add_official_header(doc, metadata, doc_type)
    else:
        display_title = _clean_title(title)
        _apply_theme(display_title)
        if display_title:
            para = doc.add_paragraph()
            run = para.add_run(display_title)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = TITLE_COLOR
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.space_after = Pt(12)
 
    sections = parse_sections(clean_text(body))
    assignments, unplaced_idx = _assign_charts_to_sections(sections, chart_images)
 
    for sec_idx, sec in enumerate(sections):
        level, heading, content, table = sec['level'], sec['heading'], sec['body'], sec['table']
 
        if level == 1:
            _add_heading(doc, heading, 15, TITLE_COLOR)
        elif level == 2:
            _add_heading(doc, heading, 13, TITLE_COLOR)
        elif level == 3:
            _add_heading(doc, heading, 11, BODY_COLOR)
 
        # Strip markdown table if present to render clean text followed by table
        tbl = extract_table(content)
        clean_content = strip_table_from_text(content) if tbl else content
 
        if clean_content:
            _add_body_content(doc, clean_content)
        if tbl:
            _add_table(doc, tbl)
 
        for idx in assignments.get(sec_idx, []):
            chart = chart_images[idx]
            img_path = chart.get("path", "") if isinstance(chart, dict) else chart[1]
            _add_image(doc, img_path)
 
    if unplaced_idx:
        para = doc.add_paragraph()
        run = para.add_run("Charts & Visualisations")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = TITLE_COLOR
        para.space_after = Pt(8)
 
        for idx in unplaced_idx:
            chart = chart_images[idx]
            img_path = chart.get("path", "") if isinstance(chart, dict) else chart[1]
            _add_image(doc, img_path)
 
    if is_official:
        _add_official_footer(doc, metadata)
 
    _add_footer(doc)
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
    doc.save(output_path)
    print(f"DOCX saved to: {output_path}")
 
 
# ── Title cleaning ─────────────────────────────────────────────────────────────
 
def _clean_title(title):
    t = re.sub(r'^[#*\-_=\s"\']+|[#*\-_=\s"\']+$', '', title).strip()
    skip = {'a', 'an', 'the', 'of', 'in', 'on', 'with', 'and', 'or', 'for', 'to', 'by'}
    words = t.split()
    result = []
    for i, w in enumerate(words):
        if i == 0 or i == len(words) - 1 or w.lower() not in skip:
            result.append(w.capitalize())
        else:
            result.append(w.lower())
    return ' '.join(result)
 
 
# ── Chart placeholder stripping ───────────────────────────────────────────────
 
def _strip_chart_placeholders(text):
    cleaned = []
    for line in text.splitlines():
        if _CHART_LINE.match(line):
            continue
        cleaned.append(line)
    return '\n'.join(cleaned)


def _clean_bracketed_placeholders(text):
    def replacer(match):
        content = match.group(1).lower().strip()
        if content.startswith('chart:') or (len(content) <= 4 and content.isupper()):
            return match.group(0)
        placeholder_words = {
            'tbd', 'insert', 'todo', 'placeholder', 'draft', 'unknown', 'xxxx', 'yyyy',
            'subject', 'ref', 'date', 'name', 'cost', 'amount', 'budget', 'value', 'price',
            'officer', 'authority', 'sign', 'signature', 'designation', 'approved', 'recommended',
            'item', 'service', 'description', 'spec', 'quantity', 'qty'
        }
        if any(w in content for w in placeholder_words) or len(content) > 4:
            if 'cost' in content or 'amount' in content or 'budget' in content or 'price' in content:
                return "Rs. __________________"
            if 'date' in content:
                return "__________________"
            return "__________________"
        return match.group(0)
    return re.sub(r'\[([^\]]+)\]', replacer, text)
 
 
# ── Deduplication ─────────────────────────────────────────────────────────────
 
def _dedupe_charts(chart_images):
    seen_paths  = set()
    seen_titles = set()
    result = []
    for chart in chart_images:
        if isinstance(chart, dict):
            ct = chart.get("title", "")
            path = chart.get("path", "")
            col = chart.get("column", "")
        else:
            ct = chart[0]
            path = chart[1]
            col = chart[2] if len(chart) > 2 else ""
            
        norm_title = ct.lower().strip()
        if path in seen_paths or norm_title in seen_titles:
            continue
        seen_paths.add(path)
        seen_titles.add(norm_title)
        
        if isinstance(chart, dict):
            result.append(chart)
        else:
            result.append({
                "title": ct,
                "path": path,
                "column": col
            })
    return result
 
 
# ── Text processing ───────────────────────────────────────────────────────────
 
def clean_text(text):
    lines, in_table = [], False
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            lines.append(line)
            continue
        if in_table:
            lines.append(line)
            if stripped == '':
                in_table = False
            continue
        m = re.match(r'^(#{1,3})\s+(.*)$', line)
        if m:
            content = re.sub(r'[*_]+', '', m.group(2)).strip()
            lines.append(f"{m.group(1)} {content}")
            continue
        if re.match(r'^[=\-_*]{3,}$', stripped):
            continue
        lines.append(re.sub(r'[*_]+', '', line))
    return "\n".join(lines)
 
 
# ── Chart-to-section matching ──────────────────────────────────────────────────
 
STOP = {'by', 'and', 'the', 'of', 'in', 'a', 'an', 'to', 'for', 'from', 'with', 'data',
        'analysis', 'report', 'file', 'product', 'wise', 'per', 'vs', 'are', 'is'}
 
def _keywords(text):
    return set(re.findall(r'\w+', text.lower())) - STOP
 
 
def _section_keywords(sec):
    heading_kw = _keywords(sec.get('heading', ''))
    table_kw = set()
    if sec.get('table'):
        for row in sec['table']:
            for cell in row:
                table_kw |= _keywords(str(cell))
    body_kw = _keywords(sec.get('body', ''))
    return heading_kw, table_kw, body_kw
 
 
def _assign_charts_to_sections(sections, chart_images):
    section_kw = [_section_keywords(sec) for sec in sections]
    assignments = {}
    unplaced = []
 
    for idx, chart in enumerate(chart_images):
        if isinstance(chart, dict):
            ct = chart.get("title", "")
            img_path = chart.get("path", "")
            col = chart.get("column", "")
        else:
            ct = chart[0]
            img_path = chart[1]
            col = chart[2] if len(chart) > 2 else ""
        
        if not os.path.exists(img_path):
            continue
 
        chart_kw = _keywords(ct)
        if col:
            chart_kw = chart_kw | _keywords(col)
            
        if not chart_kw:
            unplaced.append(idx)
            continue
 
        best_sec, best_score = None, 0
        for sec_idx, (heading_kw, table_kw, body_kw) in enumerate(section_kw):
            score = (3 * len(chart_kw & heading_kw)
                     + 2 * len(chart_kw & table_kw)
                     + 1 * len(chart_kw & body_kw))
            if score > best_score:
                best_score, best_sec = score, sec_idx
 
        if best_sec is not None and best_score > 0:
            assignments.setdefault(best_sec, []).append(idx)
        else:
            unplaced.append(idx)
 
    return assignments, unplaced
 
 
# ── Docx element builders ─────────────────────────────────────────────────────
 
def _add_heading(doc, text, size, color):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    para.space_after = Pt(8)
 
 
def _add_body_content(doc, content):
    for line in content.splitlines():
        l = line.strip()
        if not l:
            continue
        if _CHART_LINE.match(l):
            continue
        if re.match(r'^[-+•]\s+', l):
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(re.sub(r'^[-+•]\s+', '', l))
        else:
            para = doc.add_paragraph()
            run = para.add_run(l)
        run.font.size = Pt(11)
        run.font.color.rgb = BODY_COLOR
        para.space_after = Pt(4)
 
 
def _add_table(doc, table_data):
    if not table_data or len(table_data) < 2:
        return
    cols = max(len(r) for r in table_data)
    t = doc.add_table(rows=len(table_data), cols=cols)
    t.style = 'Table Grid'
 
    for ri, row in enumerate(table_data):
        while len(row) < cols:
            row.append("")
        for ci, cell_text in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = cell_text
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(10)
            if ri == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(cell, "00205b")
            elif ri % 2 == 0:
                _set_cell_bg(cell, "f2f4f8")
 
    doc.add_paragraph()
 
 
def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)
 
 
def _add_image(doc, img_path, width_inches=5.5):
    try:
        doc.add_picture(img_path, width=Inches(width_inches))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    except Exception as e:
        print(f"Could not embed image {img_path}: {e}")
 
 
def _add_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
 
    for run in para.runs:
        run.clear()
 
    run_left = para.add_run("HPGPT")
    run_left.font.size  = Pt(9)
    run_left.font.bold  = True
    run_left.font.color.rgb = FOOTER_COLOR
 
    para.add_run("\t")
 
    run_page = para.add_run()
    for tag, text in [('w:fldChar', None), ('w:instrText', 'PAGE'), ('w:fldChar', None)]:
        el = OxmlElement(tag)
        if tag == 'w:fldChar':
            el.set(qn('w:fldCharType'),
                   'begin' if not run_page._r.findall(
                       './/{%s}fldChar' % 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                   ) else 'end')
        if text:
            el.set(qn('xml:space'), 'preserve')
            el.text = text
        run_page._r.append(el)
 
    run_page.font.size  = Pt(9)
    run_page.font.bold  = True
    run_page.font.color.rgb = FOOTER_COLOR
 
    para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)